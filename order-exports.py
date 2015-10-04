from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
import csv
from collections import defaultdict as ddict

def master_transform():
	"""Transforms the CSV data"""

	# dorm_dict = ddict(list)

	with open('orders_export.csv', 'rU') as infile, \
		open('master.csv','w') as outfile1, \
		open('bad_orders.csv','w') as outfile2:
		reader = csv.reader(infile)
		reader.next()

		writer1 = csv.writer(outfile1)
		writer2 = csv.writer(outfile2)

		last_attrs = []
		for row in reader:
			# product name
			prod = row[17].lower()

			if 'refrigerator' in prod:
				prod = 'FRIDGE'
			elif 'room ready' in prod:
				prod = prod.split('-')[-1].strip().upper().split('/')[0]
			else:
				print 'Unavailable product: ', prod
				continue

			if not row[2]:
				writer1.writerow(last_attrs + [prod])
				# dorm_dict[attrs[1]].append(last_attrs + prod)
			else:
				# personal attributes
				attrs = row[45]
				attrs = attrs.upper().split('\n')

				try:
					# no need for email, reference, or netid
					attrs = [attr for attr in attrs
							if ('EMAIL' not in attr) and ('REFERENCE' not in attr) and ('NETID' not in attr)]

					name = attrs[0][9:].split()
					attrs[0] = name[0] + ' ' + name[-1]
					attrs[1] = attrs[1][10:].split('(')[-1].rstrip(')').strip() # dorm
					attrs[2] = attrs[2][12:] # room

					last_attrs = attrs

					writer1.writerow(last_attrs + [prod])

				except:
					print 'Information error: ', row[0]
					writer2.writerow(row)
					continue

				# dorm_dict[attrs[1]].append(attrs + prod)

def write_line_in_cell(cell, attr, font_size = 25):
	"""Writes info to the DOCX cell"""

	paragraph = cell.add_paragraph()
	p_format = paragraph.paragraph_format

	# no vertical spacing
	p_format.space_after = Pt(0)

	# centering
	p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# writing
	run = paragraph.add_run(attr)

	# font size
	font = run.font
	font.size = Pt(font_size)

def stickers_transform():
	"""Transforms the CSV data into a DOCX file"""
	document = Document('landscape.docx')

	with open('master9.7.csv', 'rU') as infile:
		reader = csv.reader(infile)

		table = document.add_table(2,2)

		i = 0
		for row in reader:
			index1, index2 = i / 2, i % 2
			cell = table.cell(index1, index2)

			prod = row[3]
			write_line_in_cell(cell, '\n', 20)
			write_line_in_cell(cell, row[1])
			write_line_in_cell(cell, row[2])
			write_line_in_cell(cell, row[0])
			write_line_in_cell(cell, row[3])
			write_line_in_cell(cell, '\n', 20)

			i += 1

			if i % 4 == 0:
				table = document.add_table(2,2)
				i = 0

	document.save('labels.docx')

def write_line(document, attr, font_size = 25):
	paragraph = cell.add_paragraph()
	p_format = paragraph.paragraph_format

	# no vertical spacing
	p_format.space_after = Pt(0)

	# writing
	run = paragraph.add_run(attr)

	# font size
	font = run.font
	font.size = Pt(font_size)

def checklists_transform():
	"""Transforms into checklists"""
	document = Document()

	prod_dict = ddict(list)

	with open('../orders_export 5.csv', 'rU') as infile:
		reader = csv.reader(infile)
		reader.next()

		for row in reader:
			if row[2]:
				# product name
				prod = row[17].lower()

				if 'refrigerator' in prod:
					prod = 'FRIDGE'
				elif 'room ready' in prod:
					prod = prod.split('-')[-1].strip().upper().split('/')[0]
				else:
					print 'Unavailable product: ', prod
					continue

				# personal attributes
				attrs = row[45]
				attrs = attrs.upper().split('\n')

				try:
					# no need for email, reference, or netid
					attrs = [attr for attr in attrs
							if ('EMAIL' not in attr) and ('REFERENCE' not in attr) and ('NETID' not in attr)]

					name = attrs[0][9:].split()
					attrs[0] = name[0] + ' ' + name[-1]
					attrs[1] = attrs[1][10:].split('(')[-1].rstrip(')') # dorm
					attrs[2] = attrs[2][12:] # room

					prod_dict[attrs[1] + ' ' + attrs[2]].append(attrs[0])

					# just switching up order
					attrs[0], attrs[1], attrs[2] = attrs[1], attrs[2], attrs[0]

				except:
					print 'Information error: ', row[0]
					writer.writerow(row)
					continue



if __name__ == '__main__':
	# master_transform()
	stickers_transform()